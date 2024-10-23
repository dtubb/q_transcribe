import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from PIL import Image
import typer
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from pathlib import Path
from rich import print
from rich.progress import track
import subprocess
from PyPDF2 import PdfMerger, PdfWriter, PdfReader

# Constants for page size and margins (in points: 72 points per inch)
LETTER_SIZE = (8.5 * 72, 11 * 72)  # Letter size in points
PRINTABLE_SIZE = (8.25 * 72, 10.75 * 72)  # Printable area with 0.125-inch margins

def natural_sort_key(filename: str):
    """Generate a sorting key that sorts numbers in a human-friendly way."""
    return [int(part) if part.isdigit() else part for part in re.split(r'(\d+)', filename)]

# Function to create a Word document with image on the left and transcription on the right
def create_word_document_with_layout(output_word_file: Path, image_files: list, transcription_files: list):
    doc = Document()

    # Set font for the whole document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(12)

    for image_file, transcription_file in zip(image_files, transcription_files):
        # Insert image on the left page
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.left_margin = Inches(0.125)
        section.right_margin = Inches(0.125)
        
        doc.add_picture(str(image_file), width=Inches(7.5))  # Adjust size of the image
        
        # Add new page for transcription with larger margins
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Insert the transcription text
        with open(transcription_file, "r") as f:
            transcription_text = f.read()
        p = doc.add_paragraph(transcription_text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the Word document
    doc.save(str(output_word_file))
    print(f"[green]Word document saved as {output_word_file.name}")

# Function to convert HTML to the desired formats (Word, PDF, Markdown)
def run_pandoc(input_file: Path, output_file: Path, to_format: str):
    """Run Pandoc to convert files to the specified format."""
    pandoc_command = ["pandoc", str(input_file), "-o", str(output_file), f"--to={to_format}"]
    result = subprocess.run(pandoc_command, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"[red]Error running Pandoc: {result.stderr}")
    else:
        print(f"[green]Converted {input_file.name} to {to_format.upper()} as {output_file.name}")

# Function to create image PDFs with letter size and margins
def create_image_pdf(image_file: Path, output_pdf_file: Path):
    """Convert an image to a letter-size PDF with 0.125-inch margins."""
    try:
        img = Image.open(image_file)
        img_width, img_height = img.size

        # Create a new PDF canvas with letter size (8.5 x 11 inches)
        c = canvas.Canvas(str(output_pdf_file), pagesize=letter)

        # Calculate the available space for the image with 0.125-inch margins
        max_width = 8.25 * 72  # 8.25 inches (leaving 0.125-inch margins on left and right)
        max_height = 10.75 * 72  # 10.75 inches (leaving 0.125-inch margins on top and bottom)

        # Scale the image while maintaining aspect ratio to fit the available space
        scale = min(max_width / img_width, max_height / img_height)
        scaled_width = img_width * scale
        scaled_height = img_height * scale

        # Center the image on the page
        x = (letter[0] - scaled_width) / 2  # X-coordinate for centering
        y = (letter[1] - scaled_height) / 2  # Y-coordinate for centering

        # Draw the image on the PDF canvas
        c.drawImage(str(image_file), x, y, width=scaled_width, height=scaled_height)

        # Finalize and save the PDF
        c.showPage()
        c.save()

        print(f"[green]Created image PDF: {output_pdf_file.name}")

    except Exception as e:
        print(f"[red]Error creating image PDF: {e}")

# Function to combine PDFs with image on the left and transcription on the right
from PyPDF2 import PdfMerger, PdfWriter, PdfReader

def combine_pdf_docs(output_pdf_file: Path, image_pdfs: list, transcription_pdfs: list):
    writer = PdfWriter()

    if not image_pdfs or not transcription_pdfs:
        print(f"[red]No image or transcription PDFs found to combine.")
        return

    try:
        # Add a blank title page (or replace with a custom title page)
        writer.add_blank_page(width=LETTER_SIZE[0], height=LETTER_SIZE[1])

        for image_pdf, transcription_pdf in zip(image_pdfs, transcription_pdfs):
            # Read image PDF and transcription PDF
            try:
                image_reader = PdfReader(str(image_pdf))
                transcription_reader = PdfReader(str(transcription_pdf))

                # Add all pages from the image PDF (left-hand page)
                for page in image_reader.pages:
                    writer.add_page(page)

                # Add all pages from the transcription PDF (right-hand page)
                for page in transcription_reader.pages:
                    writer.add_page(page)

                # If transcription has more pages, add blank pages to align
                if len(transcription_reader.pages) > len(image_reader.pages):
                    for _ in range(len(transcription_reader.pages) - len(image_reader.pages)):
                        writer.add_blank_page(width=LETTER_SIZE[0], height=LETTER_SIZE[1])

            except Exception as e:
                print(f"[red]Error reading PDFs: {e}")
                continue

        # Write the final merged PDF
        with open(output_pdf_file, "wb") as f_out:
            writer.write(f_out)

        print(f"[green]Combined PDF document saved as {output_pdf_file.name}")

    except Exception as e:
        print(f"[red]Error during PDF combination: {e}")

def transcribe_html_recursive(
    folder: Path = typer.Argument(..., help="Folder containing HTML or Markdown files"),
    combine_pdf: bool = typer.Option(False, help="Option to combine all PDF files into one document"),
    combine_word: bool = typer.Option(False, help="Option to combine all Word files into one document"),
    combine_markdown: bool = typer.Option(False, help="Option to combine all Markdown files into one document")
):
    print(f"[green]Processing files in {folder}")
    
    # HTML to Word, PDF, and Markdown
    for html_file in track(folder.glob("*.html"), description="Converting HTML files..."):
        word_file = html_file.with_suffix(".docx")
        pdf_file = html_file.with_suffix(".pdf")
        md_file = html_file.with_suffix(".md")
        
        # Skip conversion if the output files already exist
        if word_file.exists():
            print(f"[yellow]Skipping {word_file.name}, already exists")
        else:
            run_pandoc(html_file, word_file, "docx")

        if pdf_file.exists():
            print(f"[yellow]Skipping {pdf_file.name}, already exists")
        else:
            run_pandoc(html_file, pdf_file, "pdf")

        if md_file.exists():
            print(f"[yellow]Skipping {md_file.name}, already exists")
        else:
            run_pandoc(html_file, md_file, "markdown")

    # Generate image PDFs
    image_files = list(folder.glob("*.png")) + list(folder.glob("*.jpg")) + list(folder.glob("*.jpeg"))
    
    for image_file in track(image_files, description="Creating image PDFs..."):
        image_pdf_file = image_file.with_stem(image_file.stem + "_image").with_suffix(".pdf")
        if not image_pdf_file.exists():
            create_image_pdf(image_file, image_pdf_file)
    
    # Combine PDF documents (optional)
    if combine_pdf:
        output_pdf_file = folder / "combined_output.pdf"
        image_pdfs = sorted(folder.glob("*_image.pdf"), key=lambda f: natural_sort_key(f.stem))
        
        # Use the existing PDF files as transcriptions
        transcription_pdfs = sorted([f for f in folder.glob("*.pdf") if "_image" not in f.stem], key=lambda f: natural_sort_key(f.stem))

        if not image_pdfs or not transcription_pdfs:
            print(f"[red]No image or transcription PDFs found to combine.")
        else:
            combine_pdf_docs(output_pdf_file, image_pdfs, transcription_pdfs)

    # Combine Word documents (optional)
    if combine_word:
        output_word_file = folder / "combined_output.docx"
        image_files = sorted(folder.glob("*.png") + folder.glob("*.jpg") + folder.glob("*.jpeg"), key=lambda f: natural_sort_key(f.stem))
        transcription_files = sorted(folder.glob("*.md"), key=lambda f: natural_sort_key(f.stem))
        create_word_document_with_layout(output_word_file, image_files, transcription_files)



if __name__ == "__main__":
    typer.run(transcribe_html_recursive)
